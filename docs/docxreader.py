from docx import Document

doc = Document('demo.docx')

# get paragraph objects
p_obj = doc.paragraphs

# get paragraph text
first = doc.paragraphs[0].text
# print(first)

# runs => changes in styles
second = doc.paragraphs[1]
print(second.runs)
print(second.runs[0].text)
print(second.runs[1].text)

# boolean properties
second.runs[1].bold
second.runs[2].italic
second.runs[0].underline

# set runs to specific styles
second.runs[3].underline = True
second.runs[3].text = 'italic and underline'


# change style
second.style = 'Title'

# add paragraph
doc.add_paragraph('This paragraph is created programatically.')

# add run
second.add_run('This is a new run')
second.runs[1].bold = True

# save document
doc.save('newdoc.docx')


